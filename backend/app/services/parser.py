"""Document parsing.

Takes raw bytes + filename/mime, returns a list of ParsedPage objects with
text content and 1-indexed page numbers.

Supported formats:
  - PDF   (application/pdf)            — page numbers preserved natively
  - DOCX  (application/vnd...)         — collapsed to a single page (1)
  - TXT   (text/plain)                 — single page
  - MD    (text/markdown)              — single page

Why DOCX is one page: the .docx format is paragraph-flow; "pages" depend on
print rendering and aren't reliably retrievable from the file. Treating it
as a single page is the honest answer; citations on DOCX content read
"page 1 of <filename>".
"""

from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.errors import IngestionError
from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass(frozen=True, slots=True)
class ParsedPage:
    """One page of a parsed document. `page` is 1-indexed."""

    page: int
    text: str


# MIME type -> internal format key
MIME_TO_FORMAT: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/x-markdown": "md",
}

# Filename extension fallback when MIME is unreliable (browsers often send
# octet-stream or wrong types for less-common extensions).
EXT_TO_FORMAT: dict[str, str] = {
    "pdf": "pdf",
    "docx": "docx",
    "txt": "txt",
    "md": "md",
    "markdown": "md",
}


def parse(filename: str, content: bytes, content_type: str | None = None) -> list[ParsedPage]:
    """Parse a binary document into a list of ParsedPage.

    Format is detected from `content_type` first, then falls back to the
    filename extension. Empty pages are dropped. Returns an empty list if
    the document parsed but contained no extractable text — callers should
    treat that as a failure.

    Raises:
        IngestionError: format unsupported, file unreadable, or PDF encrypted.
    """
    fmt = _detect_format(filename, content_type)
    logger.info("parser_dispatch", filename=filename, format=fmt, size=len(content))

    if fmt == "pdf":
        return _parse_pdf(content)
    if fmt == "docx":
        return _parse_docx(content)
    if fmt in ("txt", "md"):
        return _parse_text(content)

    # Should be unreachable thanks to _detect_format, kept for safety.
    raise IngestionError(f"unsupported format: {fmt!r}", code="unsupported_format")


def _detect_format(filename: str, content_type: str | None) -> str:
    if content_type:
        # Strip any "; charset=..." parameters before lookup.
        primary = content_type.split(";", 1)[0].strip().lower()
        if primary in MIME_TO_FORMAT:
            return MIME_TO_FORMAT[primary]

    if "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext in EXT_TO_FORMAT:
            return EXT_TO_FORMAT[ext]

    raise IngestionError(
        f"could not determine format from filename={filename!r} content_type={content_type!r}",
        code="unsupported_format",
    )


def _parse_pdf(content: bytes) -> list[ParsedPage]:
    try:
        reader = PdfReader(io.BytesIO(content))
    except PdfReadError as e:
        raise IngestionError(f"failed to parse PDF: {e}", code="pdf_invalid") from e
    except Exception as e:
        raise IngestionError(f"failed to open PDF: {e}", code="pdf_invalid") from e

    if reader.is_encrypted:
        raise IngestionError(
            "PDF is encrypted; please decrypt before uploading",
            code="pdf_encrypted",
        )

    pages: list[ParsedPage] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001 - defensive; corrupt pages
            logger.warning("pdf_page_extract_failed", page=index, error=str(e))
            raw = ""
        text = _normalize_whitespace(raw)
        if text:
            pages.append(ParsedPage(page=index, text=text))
    return pages


def _parse_docx(content: bytes) -> list[ParsedPage]:
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as e:
        raise IngestionError(f"failed to parse DOCX: {e}", code="docx_invalid") from e

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Tables often hold structured policy info — pull cell text as separate "lines".
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    full = _normalize_whitespace("\n".join(parts))
    if not full:
        return []
    return [ParsedPage(page=1, text=full)]


def _parse_text(content: bytes) -> list[ParsedPage]:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        # Many internal docs are latin-1 / windows-1252; replace bad bytes
        # rather than fail outright.
        text = content.decode("latin-1", errors="replace")
        logger.warning("text_fallback_decoding", encoding="latin-1")

    text = _normalize_whitespace(text)
    if not text:
        return []
    return [ParsedPage(page=1, text=text)]


def _normalize_whitespace(text: str) -> str:
    """Trim trailing spaces and collapse runs of blank lines.

    Keeps single blank lines as paragraph delimiters so chunkers can use
    them as natural split points later.
    """
    out: list[str] = []
    blank_run = 0
    for line in text.splitlines():
        stripped = line.rstrip()
        if not stripped:
            blank_run += 1
            if blank_run == 1 and out:
                out.append("")
            continue
        blank_run = 0
        out.append(stripped)
    return "\n".join(out).strip()
