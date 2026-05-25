"""Shared pytest fixtures."""

from __future__ import annotations

from collections.abc import AsyncIterator
from io import BytesIO

import pytest
import pytest_asyncio
from fastapi.testclient import TestClient
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.adapters.db import SessionLocal, get_db_session
from app.main import app


@pytest.fixture
def client() -> TestClient:
    """Synchronous test client. FastAPI's TestClient drives the ASGI app
    in-process, so no real network is involved."""
    return TestClient(app)


# --- Async / DB integration fixtures -------------------------------------
# These tie a test's HTTP requests to a real DB session that rolls back
# at teardown, giving every test a clean slate without a separate test DB.


@pytest_asyncio.fixture
async def db() -> AsyncIterator[AsyncSession]:
    """A real DB session, rolled back at end of test. Requires a running db."""
    async with SessionLocal() as session:
        try:
            yield session
        finally:
            await session.rollback()


@pytest_asyncio.fixture
async def async_client(db: AsyncSession) -> AsyncIterator[AsyncClient]:
    """HTTP client whose requests share the test's DB session.

    Routes call get_db_session via Depends; we override it to yield the
    same session the test fixture is using so writes are visible to the
    test, and rolled back at teardown.
    """

    async def _override() -> AsyncIterator[AsyncSession]:
        yield db  # do NOT commit; the fixture owns the lifecycle

    app.dependency_overrides[get_db_session] = _override
    try:
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as ac:
            yield ac
    finally:
        app.dependency_overrides.clear()


# --- Document fixtures ----------------------------------------------------
# We generate fixture files at test time instead of committing binaries.
# Each fixture returns (filename, content_bytes) so tests can drive the
# parser exactly as it would be called from the upload endpoint.


@pytest.fixture
def sample_txt() -> tuple[str, bytes]:
    body = (
        "Employee handbook excerpt.\n"
        "All employees are entitled to fifteen (15) PTO days per year.\n"
        "\n"
        "Carry-over is capped at five (5) days into the next calendar year."
    )
    return "handbook.txt", body.encode("utf-8")


@pytest.fixture
def sample_md() -> tuple[str, bytes]:
    body = (
        "# Information Security Policy\n"
        "\n"
        "## Passwords\n"
        "All employee passwords must be at least 12 characters long.\n"
        "\n"
        "## MFA\n"
        "Multi-factor authentication is **required** for production access."
    )
    return "security_policy.md", body.encode("utf-8")


@pytest.fixture
def sample_docx() -> tuple[str, bytes]:
    """A small DOCX with a heading, paragraphs, and a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("IT Helpdesk SOP", level=1)
    doc.add_paragraph("This document outlines the standard helpdesk response process.")
    doc.add_paragraph("Tickets are triaged within one business hour of receipt.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Severity"
    table.rows[0].cells[1].text = "Response time"
    table.rows[1].cells[0].text = "Critical"
    table.rows[1].cells[1].text = "15 minutes"

    buf = BytesIO()
    doc.save(buf)
    return "helpdesk_sop.docx", buf.getvalue()


@pytest.fixture
def sample_pdf_two_pages() -> tuple[str, bytes]:
    """A 2-page PDF generated with reportlab (test-only dependency)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)

    # Page 1
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, "Employee Benefits Summary")
    c.drawString(72, 690, "Health insurance is fully employer-paid for full-time employees.")
    c.drawString(72, 670, "Dental and vision coverage is available at a 50% subsidy.")
    c.showPage()

    # Page 2
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, "Retirement Plans")
    c.drawString(72, 690, "401(k) matching up to 4% of base salary, vesting over four years.")
    c.drawString(72, 670, "Employees may enroll on the first of the month after hire.")
    c.save()

    return "benefits.pdf", buf.getvalue()
